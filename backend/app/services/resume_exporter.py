import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def generate_resume_docx(optimized_resume: dict) -> bytes:
    """
    Generates a .docx file from the optimized_resume JSON object.
    Returns the file as raw bytes (to be used in a BytesIO stream).
    """
    doc = Document()
    
    # Adjust default margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 0. NAME (Large and Bold)
    if optimized_resume.get("name"):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(optimized_resume["name"].upper())
        run.bold = True
        run.font.size = Pt(16)
        p.paragraph_format.space_after = Pt(4)

    # 0.1 CONTACT INFO (Centered, smaller)
    if optimized_resume.get("contact_info"):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(optimized_resume["contact_info"])
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(12)

    # 1. OBJECTIVE (If present)
    if optimized_resume.get("objective"):
        p = doc.add_paragraph()
        run = p.add_run(optimized_resume["objective"])
        run.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(12)

    # 2. Summary
    if optimized_resume.get("summary"):
        p = doc.add_paragraph()
        run = p.add_run(optimized_resume["summary"])
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(12)

    # 3. Experience
    if optimized_resume.get("experience"):
        doc.add_heading("Experience", level=1)
        for exp in optimized_resume["experience"]:
            # Role and Company
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.get('role', '')} at {exp.get('company', '')}")
            run.bold = True
            run.font.size = Pt(11)
            
            # Dates
            if exp.get("dates"):
                p = doc.add_paragraph()
                run = p.add_run(exp["dates"])
                run.font.size = Pt(10)
                run.italic = True
                p.paragraph_format.space_after = Pt(4)

            # Bullets
            for bullet_text in exp.get("bullets", []):
                p = doc.add_paragraph(bullet_text, style='List Bullet')
                for run in p.runs:
                    run.font.size = Pt(10)
            doc.add_paragraph() # Spacer

    # 4. Skills
    if optimized_resume.get("skills"):
        doc.add_heading("Skills", level=1)
        for skill_group in optimized_resume["skills"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{skill_group.get('category', '')}: ")
            run.bold = True
            run.font.size = Pt(11)
            
            skills_text = ", ".join(skill_group.get("skills", []))
            run = p.add_run(skills_text)
            run.font.size = Pt(11)

    # 5. Education
    if optimized_resume.get("education"):
        doc.add_heading("Education", level=1)
        for edu in optimized_resume["education"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu.get('degree', '')} - {edu.get('institution', '')}")
            run.bold = True
            run.font.size = Pt(11)
            
            if edu.get("dates"):
                p = doc.add_paragraph()
                run = p.add_run(edu["dates"])
                run.font.size = Pt(10)
                run.italic = True

    # Save document to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream.getvalue()