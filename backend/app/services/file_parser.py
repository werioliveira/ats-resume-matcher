import io
import re
from collections import Counter

import pdfplumber
from docx import Document


class UnextractableFileError(Exception):
    """Raised when PDF is scanned, image-based, or has no extractable text."""
    pass


# Keywords to help identify section headers in resumes (EN and PT)
HEADER_KEYWORDS = {
    "experience", "education", "skills", "summary", "objective",
    "projects", "certifications", "languages", "contact", "profile",
    "work history", "professional experience", "academic background",
    "experiência", "educação", "habilidades", "resumo", "projetos",
    "certificações", "idiomas", "contato", "perfil", "formação",
    "histórico profissional", "experiência profissional"
}

# Tolerance in points to group characters/words into the same line
Y_TOLERANCE = 3


def parse_resume(file_bytes: bytes, content_type: str) -> str:
    """
    Extract text from PDF or DOCX resume.
    Returns Markdown-like structured text.
    """
    if content_type == "application/pdf":
        return _parse_pdf(file_bytes)
    elif content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ):
        return _parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {content_type}")


def _parse_pdf(file_bytes: bytes) -> str:
    """Extracts text from PDF, handling 2-column layouts and header inference."""
    all_text_parts = []
    all_words = []

    # Wrap in try/except to catch corrupted or empty files before parsing
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                # REMOVED "bold" from extra_attrs to prevent KeyError on most PDFs
                words = page.extract_words(extra_attrs=["size", "fontname"])
                
                if not words:
                    text = page.extract_text()
                    if text:
                        all_text_parts.append(text)
                    continue

                all_words.extend(words)

                if _is_two_column_layout(words, page.width):
                    page_text = _extract_two_column(words)
                else:
                    page_text = _extract_single_column(words)

                all_text_parts.append(page_text)
    except Exception as e:
        # Catches pdfminer errors for empty bytes, invalid headers, etc.
        raise UnextractableFileError(
            f"Failed to read PDF file. It may be corrupted or invalid. Details: {str(e)}"
        )

    raw_text = "\n".join(all_text_parts)

    if not raw_text.strip():
        raise UnextractableFileError(
            "No text could be extracted from this PDF. It might be scanned or image-based."
        )

    raw_text = _infer_pdf_headers(raw_text, all_words)
    
    return _normalize_text(raw_text)


def _is_two_column_layout(words: list[dict], page_width: float) -> bool:
    """Heuristic to detect if a page uses a 2-column layout."""
    if not words:
        return False
    
    x_positions = [w["x0"] for w in words]
    mid_point = page_width / 2
    
    left_count = sum(1 for x in x_positions if x < mid_point)
    right_count = sum(1 for x in x_positions if x >= mid_point)
    
    total = left_count + right_count
    if total == 0:
        return False
    
    left_ratio = left_count / total
    # If both sides have significant text, it's 2-column
    return 0.20 < left_ratio < 0.80


def _extract_two_column(words: list[dict]) -> str:
    """Extracts text reading left-to-right, top-to-bottom across columns."""
    lines = _group_words_into_lines(words)
    return "\n".join(lines)


def _extract_single_column(words: list[dict]) -> str:
    """Extracts text reading top-to-bottom."""
    lines = _group_words_into_lines(words)
    return "\n".join(lines)


def _group_words_into_lines(words: list[dict]) -> list[str]:
    """Groups words by approximate Y coordinate, sorts X within line, joins."""
    rows = {}
    for w in words:
        y_key = round(w["top"] / Y_TOLERANCE) * Y_TOLERANCE
        if y_key not in rows:
            rows[y_key] = []
        rows[y_key].append(w)

    sorted_lines = []
    for y in sorted(rows.keys()):
        # Sort words in the row left to right
        row_words = sorted(rows[y], key=lambda w: w["x0"])
        
        # extract_words already gives us complete words (e.g., "São", "Gabriel").
        # We just need to join them with a space. Any extra spaces are cleaned
        # up later by the _normalize_text() function.
        line_text = " ".join(w["text"] for w in row_words)
        
        if line_text.strip():
            sorted_lines.append(line_text.strip())
            
    return sorted_lines


def _infer_pdf_headers(text: str, all_words: list[dict]) -> str:
    """Replaces lines that had larger/bold fonts and match keywords with Markdown headers."""
    if not all_words:
        return text

    sizes = [round(w.get("size", 12), 1) for w in all_words]
    if not sizes:
        return text
        
    most_common_size = Counter(sizes).most_common(1)[0][0]

    line_props = {}
    for w in all_words:
        y_key = round(w["top"] / Y_TOLERANCE) * Y_TOLERANCE
        if y_key not in line_props:
            line_props[y_key] = {"sizes": [], "is_bold": []}
        line_props[y_key]["sizes"].append(w.get("size", 12))
        # Detect bold by checking if "bold" is in the font name (e.g., "Arial-BoldMT")
        fontname = w.get("fontname", "").lower()
        line_props[y_key]["is_bold"].append("bold" in fontname)

    lines = text.split("\n")
    result_lines = []
    
    sorted_y_keys = sorted(line_props.keys())
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            result_lines.append(line)
            continue

        if i < len(sorted_y_keys):
            y_key = sorted_y_keys[i]
            props = line_props[y_key]
            avg_size = sum(props["sizes"]) / len(props["sizes"]) if props["sizes"] else most_common_size
            has_bold = any(props["is_bold"])
        else:
            avg_size = most_common_size
            has_bold = False

        is_keyword = any(kw in stripped.lower() for kw in HEADER_KEYWORDS)
        
        if is_keyword and (avg_size > most_common_size * 1.05 or has_bold):
            result_lines.append(f"## {stripped}")
        elif avg_size > most_common_size * 1.2 and is_keyword:
             result_lines.append(f"## {stripped}")
        else:
            result_lines.append(line)

    return "\n".join(result_lines)


def _parse_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX, identifying styles and tables."""
    doc = Document(io.BytesIO(file_bytes))
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect headings by style name
        if para.style and para.style.name and "heading" in para.style.name.lower():
            lines.append(f"## {text}")
        # Detect bold paragraphs as likely section headers
        elif para.runs and all(run.bold for run in para.runs if run.text.strip()):
            lines.append(f"## {text}")
        else:
            lines.append(text)

    # Extract tables (common in older or template-based resumes)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)

    raw_text = "\n".join(lines)
    return _normalize_text(raw_text)


def _normalize_text(text: str) -> str:
    """Cleans up extracted text to consistent Markdown-like format."""
    # Remove excessive blank lines (more than 2 consecutive)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove excessive spaces within lines
    text = re.sub(r"[^\S\n]{2,}", " ", text)
    
    # FIX: Add spaces around smashed symbols (e.g., "Testes&Ferramentas" -> "Testes & Ferramentas")
    text = re.sub(r'([a-zA-Z0-9])([&|/])', r'\1 \2 ', text)
    text = re.sub(r'([&|/])([a-zA-Z0-9])', r' \1 \2', text)
    
    # Remove trailing spaces per line
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return text.strip()